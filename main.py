from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
import shutil
from datetime import datetime
from anthropic import Anthropic
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize Anthropic client
client = Anthropic(
    api_key=os.environ.get("ANTHROPIC_API_KEY"),
)

def create_output_folder(folder_name=None):
    """Create an output folder for saving files."""
    if folder_name is None:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        folder_name = f"output_{timestamp}"

    output_folder = os.path.join(os.path.dirname(__file__), folder_name)

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Created output folder: {output_folder}")
    else:
        print(f"Using existing output folder: {output_folder}")

    return output_folder

def merge_docx_files(skills_path, job_desc_path, output_folder, output_filename=None):
    """
    Loads two .docx files, creates a copy of the first one,
    and appends the second one to the first.
    """
    # Create output path if not provided
    if output_filename is None:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output_filename = f"skills_prompt_w_job-{timestamp}.docx"
        # Save to the current directory, not in the assets folder
    output_path = os.path.join(output_folder, output_filename)

    print('Creating a copy of the skills prompt file')
    skills_doc = Document(skills_path)

    print('Loading the job description file.')
    job_desc_doc = Document(job_desc_path)

    skills_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    skills_doc.add_heading('Job Description', level=1)

    for paragraph in job_desc_doc.paragraphs: 
        new_paragraph = skills_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

    skills_doc.save(output_path)
    print(f"Merged document saved to: {output_path}")
    
    return output_path

def extract_text_from_docx(docx_path):
    """
    Extract all text from a .docx file
    """
    doc = Document(docx_path)
    full_text = []
    
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    return "\n".join(full_text)

def ask_claude(prompt):
    """
    Send a prompt to Claude API and get the response.
    Uses Claude 3.7 Sonnet with increased token limit.
    """
    print("Sending prompt to Claude API...")
    
    message = client.messages.create(
        max_tokens=8192,  # Increased from 4096 to 8192
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="claude-3-7-sonnet-20250219",  # Updated to 3.7 Sonnet
    )
    
    return message.content[0].text

def save_response_to_docx(response_text, output_folder, output_filename=None):
    """
    Save Claude's response to a new .docx file
    """
    if output_filename is None:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output_filename = f"claude_response-{timestamp}.docx"
    
    output_path = os.path.join(output_folder, output_filename)

    doc = Document()
    doc.add_heading('Claude Response', level=1)
    
    # Split the response by paragraphs and add each to the document
    paragraphs = response_text.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty paragraphs
            doc.add_paragraph(para)
    
    doc.save(output_path)
    print(f"Claude's response saved to: {output_path}")
    
    return output_path

def extract_skills_from_response(docx_path):
    """
    Extract main skills and sub-skills from Claude's response.
    """
    text = extract_text_from_docx(docx_path)

    lines = [line.strip() for line in text.split('\n') if line.strip()]

    if lines and lines[0].startswith('Claude Response'):
        lines = lines[1:]
    
    main_item_pattern = re.compile(r'^(\d+)\)\s+(.*)$')
    sub_item_pattern = re.compile(r'^\(([ivxl]+)\)\s+(.*)$')

    main_items_text = []
    sub_items_text = []

    for line in lines:
        main_match = main_item_pattern.match(line)
        if main_match:
            text = main_match.group(2)
            main_items_text.append(text)
            continue

        sub_match = sub_item_pattern.match(line)
        if sub_match:
            text = sub_match.group(2)
            sub_items_text.append(text)
    
    return main_items_text, sub_items_text

def update_skills_table(cv_template_path, main_skills, sub_skills, output_folder, output_filename=None):
    """
    Updates the existing empty skills table in the CV template with the provided skills,
    with all text left-aligned.
    
    Parameters:
    cv_template_path (str): Path to the CV template document with existing table
    main_skills (list): List of main skills
    sub_skills (list): List of sub-skills (5 sub-skills per main skill)
    output_folder (str): Folder to save the output document
    output_filename (str, optional): Name for the output file
    
    Returns:
    str: Path to the saved document
    """
    
    # Create output filename if not provided
    if output_filename is None:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output_filename = f"CV_with_skills-{timestamp}.docx"
    
    output_path = os.path.join(output_folder, output_filename)
    
    # Make a copy of the template first to avoid modifying the original
    temp_path = os.path.join(output_folder, "temp_" + os.path.basename(cv_template_path))
    shutil.copy2(cv_template_path, temp_path)
    
    # Load the CV template copy
    doc = Document(temp_path)
    
    # Find the skills table (first table in the document)
    if len(doc.tables) == 0:
        print("Warning: No tables found in the template.")
        os.remove(temp_path)  # Clean up
        return None
    
    # Get the first table
    skills_table = doc.tables[0]
    
    # Make sure we have the right number of sub-skills (5 per main skill)
    if len(sub_skills) < len(main_skills) * 5:
        print(f"Warning: Not enough sub-skills provided. Expected {len(main_skills) * 5}, got {len(sub_skills)}")
        # Pad with empty strings if necessary
        sub_skills.extend([''] * (len(main_skills) * 5 - len(sub_skills)))
    
    # Calculate how many skill pairs we need (2 columns per row)
    num_pairs = (len(main_skills) + 1) // 2  # Ceiling division to handle odd number of main skills
    
    # Add more rows if needed
    while len(skills_table.rows) < num_pairs:
        skills_table.add_row()  # Add more rows if needed
    
    # Fill the table with skills
    for pair_idx in range(num_pairs):
        # Calculate indices for this pair
        left_skill_idx = pair_idx * 2
        right_skill_idx = pair_idx * 2 + 1
        
        # Get the current row
        row = skills_table.rows[pair_idx]
        
        # Update left cell
        if left_skill_idx < len(main_skills):
            left_cell = row.cells[0]
            
            # Check if the cell has any paragraphs
            if len(left_cell.paragraphs) == 0:
                left_cell.add_paragraph()
            
            # Clear any existing content
            for i in range(len(left_cell.paragraphs)):
                if i > 0:  # Keep the first paragraph, remove extras
                    left_cell._element.remove(left_cell.paragraphs[i]._p)
            
            # Now we have just one empty paragraph
            p = left_cell.paragraphs[0]
            p.text = ""  # Ensure it's empty
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment
            run = p.add_run(main_skills[left_skill_idx])
            run.bold = True
            
            # Add sub-skills for left skill
            for i in range(5):
                sub_idx = left_skill_idx * 5 + i
                if sub_idx < len(sub_skills) and sub_skills[sub_idx]:
                    p = left_cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment
                    p.add_run(sub_skills[sub_idx])
        
        # Update right cell
        if right_skill_idx < len(main_skills):
            right_cell = row.cells[1]
            
            # Check if the cell has any paragraphs
            if len(right_cell.paragraphs) == 0:
                right_cell.add_paragraph()
            
            # Clear any existing content
            for i in range(len(right_cell.paragraphs)):
                if i > 0:  # Keep the first paragraph, remove extras
                    right_cell._element.remove(right_cell.paragraphs[i]._p)
            
            # Now we have just one empty paragraph
            p = right_cell.paragraphs[0]
            p.text = ""  # Ensure it's empty
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment
            run = p.add_run(main_skills[right_skill_idx])
            run.bold = True
            
            # Add sub-skills for right skill
            for i in range(5):
                sub_idx = right_skill_idx * 5 + i
                if sub_idx < len(sub_skills) and sub_skills[sub_idx]:
                    p = right_cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment
                    p.add_run(sub_skills[sub_idx])
    
    # Save the document
    doc.save(output_path)
    print(f"CV with updated skills saved to: {output_path}")
    
    # Clean up the temporary file
    os.remove(temp_path)
    
    return output_path

def extract_summary_from_response(response_text):
    """
    Extract the professional summary from Claude's response text.
    
    Parameters:
    response_text (str): Claude's response text
    
    Returns:
    str: Extracted professional summary
    """
    # This is a simple extraction that assumes the first paragraph is the summary
    # You might want to customize this based on Claude's actual response format
    
    # Strip any leading/trailing whitespace and split by new lines
    lines = [line.strip() for line in response_text.split('\n') if line.strip()]
    
    # Skip any lines that might be headers or instructions
    # For example, if Claude responds with "Here's a professional summary:" first
    start_idx = 0
    for i, line in enumerate(lines):
        if "summary" in line.lower() and ":" in line:
            start_idx = i + 1
            break
    
    # Take the first substantive paragraph after any headers
    if start_idx < len(lines):
        return lines[start_idx]
    
    # Fallback: just return the first non-empty line
    for line in lines:
        if line and not line.startswith("#") and not line.startswith("Here"):
            return line
    
    return ""

def add_professional_summary(cv_path, summary_text, output_folder, output_filename=None):
    """
    Adds a professional summary paragraph below the "Professional Summary" header in the CV
    while preserving all formatting and document structure.
    
    Parameters:
    cv_path (str): Path to the CV document
    summary_text (str): The professional summary text to add
    output_folder (str): Folder to save the output document
    output_filename (str, optional): Name for the output file
    
    Returns:
    str: Path to the saved document
    """
    # Create output filename if not provided
    if output_filename is None:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output_filename = f"CV_with_summary-{timestamp}.docx"
    
    output_path = os.path.join(output_folder, output_filename)
    
    # Make a copy of the CV first to avoid modifying the original
    temp_path = os.path.join(output_folder, "temp_" + os.path.basename(cv_path))
    shutil.copy2(cv_path, temp_path)
    
    # Load the CV copy
    doc = Document(temp_path)
    
    # Find the "Professional Summary" section
    summary_section_found = False
    summary_section_index = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        if "Professional Summary" in paragraph.text:
            summary_section_found = True
            summary_section_index = i
            break
    
    if not summary_section_found:
        print("Warning: 'Professional Summary' section not found in the CV.")
        os.remove(temp_path)  # Clean up
        return None
    
    # Simple approach: Simply add a new paragraph right after the Professional Summary header
    if summary_section_index is not None:
        # Get the summary paragraph's parent element
        header_paragraph = doc.paragraphs[summary_section_index]
        
        # Insert new paragraph after the header
        new_para = doc.add_paragraph()
        new_para.text = summary_text
        
        # This is the tricky part - we need to make sure the new paragraph appears
        # right after the Professional Summary header
        # We'll use the xml element directly
        body_element = header_paragraph._element.getparent()
        
        # Insert the new paragraph right after the header
        header_index = body_element.index(header_paragraph._element)
        body_element.insert(header_index + 1, new_para._element)
    
    # Save the document
    doc.save(output_path)
    print(f"CV with professional summary saved to: {output_path}")
    
    # Clean up the temporary file
    os.remove(temp_path)
    
    return output_path

def add_cover_letter_content(template_path, cover_letter_text, output_folder, output_filename="2025_Cover_Letter.docx"):
    """
    Adds cover letter content from Claude's response to the cover letter template.
    The content is added above "Thank you for your consideration,"
    
    Parameters:
    template_path (str): Path to the cover letter template document
    cover_letter_text (str): The cover letter content to add
    output_folder (str): Folder to save the output document
    output_filename (str): Name for the output file (default: "2025_Cover_Letter.docx")
    
    Returns:
    str: Path to the saved document
    """
    
    output_path = os.path.join(output_folder, output_filename)
    
    # Make a copy of the template first to avoid modifying the original
    temp_path = os.path.join(output_folder, "temp_" + os.path.basename(template_path))
    shutil.copy2(template_path, temp_path)
    
    # Load the cover letter template copy
    doc = Document(temp_path)
    
    # Find the "Thank you for your consideration" paragraph
    thank_you_index = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        if "Thank you for your consideration" in paragraph.text:
            thank_you_index = i
            break
    
    if thank_you_index is None:
        print("Warning: 'Thank you for your consideration' text not found in the template.")
        os.remove(temp_path)  # Clean up
        return None
    
    # Extract company name for greeting if available
    company_name = ""
    greeting = f"Dear {company_name} Hiring Manager," if company_name else "Dear Hiring Manager,"
    
    # Process the cover letter text - make sure there's a blank line after greeting
    # and before the content from Claude's response
    paragraphs = [greeting, ""] + cover_letter_text.split('\n')
    
    # Insert the cover letter content before the "Thank you" paragraph
    # We need to insert in reverse order to maintain correct ordering
    for i, para_text in enumerate(reversed(paragraphs)):
        if para_text.strip():  # Only add non-empty paragraphs
            # Get the "Thank you" paragraph's parent element
            thank_you_paragraph = doc.paragraphs[thank_you_index]
            body_element = thank_you_paragraph._element.getparent()
            
            # Create a new paragraph for this line of text
            new_para = doc.add_paragraph(para_text)
            
            # Insert the new paragraph before the "Thank you" paragraph
            body_element.insert(body_element.index(thank_you_paragraph._element), new_para._element)
    
    # Add a blank line between the cover letter content and "Thank you"
    thank_you_paragraph = doc.paragraphs[thank_you_index]
    body_element = thank_you_paragraph._element.getparent()
    blank_para = doc.add_paragraph("")
    body_element.insert(body_element.index(thank_you_paragraph._element), blank_para._element)
    
    # Save the document
    doc.save(output_path)
    print(f"Cover letter saved to: {output_path}")
    
    # Clean up the temporary file
    os.remove(temp_path)
    
    return output_path

def extract_cover_letter_from_response(response_text):
    """
    Extract the cover letter content from Claude's response.
    
    Parameters:
    response_text (str): Claude's response text
    
    Returns:
    str: Extracted cover letter content
    """
    # Strip any leading/trailing whitespace and split by new lines
    lines = [line.strip() for line in response_text.split('\n') if line.strip()]
    
    # Skip any lines that might be headers or instructions
    # For example, if Claude responds with explanatory text first
    start_idx = 0
    
    # Look for common greeting patterns that indicate the start of the letter
    for i, line in enumerate(lines):
        if (line.startswith("Dear") or 
            "Hiring Manager" in line or 
            "Recruiter" in line or
            "Recruitment" in line):
            start_idx = i
            break
    
    # If we didn't find a greeting, try to find the first paragraph that looks like content
    if start_idx == 0:
        for i, line in enumerate(lines):
            # Skip lines that look like Claude's explanatory text
            if (line.startswith("Here") or 
                "cover letter" in line.lower() or 
                "draft" in line.lower() or
                line.startswith("#")):
                continue
            else:
                start_idx = i
                break
    
    # Get all content from the start to the end, omitting any closing like "Sincerely"
    end_idx = len(lines)
    for i, line in enumerate(lines[start_idx:], start_idx):
        if (line.startswith("Sincerely") or 
            line.startswith("Best") or 
            line.startswith("Regards") or
            line.startswith("Thank you")):
            end_idx = i
            break
    
    # Extract the cover letter content
    cover_letter_content = "\n".join(lines[start_idx:end_idx])
    
    return cover_letter_content

if __name__ == "__main__":
    # Create output folder
    output_folder = create_output_folder("job_application_outputs")

    # Define paths relative to the script
    assets_folder = os.path.join(os.path.dirname(__file__), "assets")
    skills_path = os.path.join(assets_folder, "skills-prompt.docx")
    job_desc_path = os.path.join(assets_folder, "job-description.docx")
    personal_summary_prompt_path = os.path.join(assets_folder, "personal_summary_prompt.docx")
    cv_template_path = os.path.join(assets_folder, "CV_Template.docx")
    cover_letter_template_path = os.path.join(assets_folder, "cover_letter_template.docx")
    
    # Verify files exist
    if not os.path.exists(skills_path):
        raise FileNotFoundError(f"Skills prompt file not found at {skills_path}")
    if not os.path.exists(job_desc_path):
        raise FileNotFoundError(f"Job description file not found at {job_desc_path}")
    if not os.path.exists(cv_template_path):
        raise FileNotFoundError(f"CV template file not found at {cv_template_path}")
    
    # Step 1: Merge the documents
    merged_file = merge_docx_files(skills_path, job_desc_path, output_folder)
    print(f"Successfully created merged document: {merged_file}")
    
    # Step 2: Extract text from the merged document
    prompt_text = extract_text_from_docx(merged_file)
    print(f"Extracted {len(prompt_text)} characters from the merged document")
    
    # Step 3: Send to Claude API
    claude_response = ask_claude(prompt_text)
    print("Received skills response from Claude")
    
    # Step 4: Save Claude's response to a new document
    response_file = save_response_to_docx(claude_response, output_folder)
    print(f"Response saved to: {response_file}")

    # Step 5: Extract skills from Claude's response
    main_items, sub_items = extract_skills_from_response(response_file)
    print(f"Extracted {len(main_items)} main skills and {len(sub_items)} sub-skills")
    
    # Step 6: Update the skills table in the CV template
    updated_cv = update_skills_table(
        cv_template_path, 
        main_items,  # Your main skills list
        sub_items,   # Your sub-skills list
        output_folder
    )
    print(f"Updated CV saved to: {updated_cv}")

    # Step 7: Merge the job description to the personal summary prompt document
    personal_summary_prompt = merge_docx_files(personal_summary_prompt_path, job_desc_path, output_folder)
    print(f"Successfully created merged document: {personal_summary_prompt}")

    # Step 8: Extract text from the personal summary prompt
    summary_prompt_text = extract_text_from_docx(personal_summary_prompt)
    print(f"Extracted {len(summary_prompt_text)} characters from the summary prompt")

    # Step 9: Ask Claude to create a personal summary paragraph
    claude_summary_response = ask_claude(summary_prompt_text)
    print("Received personal summary response from Claude")

    # Step 10: Extract the summary from Claude's response
    summary_text = extract_summary_from_response(claude_summary_response)
    print(f"Extracted summary: {summary_text[:100]}...")  # Print the first 100 chars

    # Step 11: Add the summary to the CV (use the updated CV from step 6)
    final_cv = add_professional_summary(
        updated_cv,
        summary_text,
        output_folder,
        "final_CV.docx"  # Give it a specific name
    )
    print(f"Final CV with summary and skills saved to: {final_cv}")

    # Step 12: Verify cover letter template exists
    if not os.path.exists(cover_letter_template_path):
        raise FileNotFoundError(f"Cover letter template file not found at {cover_letter_template_path}")
    
    if not os.path.exists(personal_summary_prompt_path):
        print(f"Warning: Personal summary prompt file not found at {personal_summary_prompt_path}")
        # You could create a default one here if needed
    
    # Step 13: Merge the cover letter prompt document with the relevant job description document
    cover_letter_prompt_path = os.path.join(assets_folder, "cover_letter_prompt_long.docx")
    if not os.path.exists(cover_letter_prompt_path):
        raise FileNotFoundError(f"Cover letter prompt file not found at {cover_letter_prompt_path}")
    
    cover_letter_prompt = merge_docx_files(cover_letter_prompt_path, job_desc_path, output_folder)
    print(f"Successfully created merged document for cover letter: {cover_letter_prompt}")

    # Step 14: Extract the complete cover letter prompt document text for the API call
    cover_letter_prompt_text = extract_text_from_docx(cover_letter_prompt)
    print(f"Extracted {len(cover_letter_prompt_text)} characters from the cover letter prompt")

    # Step 15: Ask Claude to create a cover letter
    claude_cover_letter_response = ask_claude(cover_letter_prompt_text)
    print("Received cover letter response from Claude")
    
    # Step 16: Save Claude's cover letter response to a document
    cover_letter_response_file = save_response_to_docx(claude_cover_letter_response, output_folder)
    print(f"Cover letter response saved to: {cover_letter_response_file}")

    # Step 17: Extract the cover letter content from Claude's response
    cover_letter_content = extract_cover_letter_from_response(claude_cover_letter_response)
    print(f"Extracted cover letter content: {cover_letter_content[:100]}...")  # Print the first 100 chars

    # Step 18: Add the cover letter content to the template
    final_cover_letter = add_cover_letter_content(
        cover_letter_template_path,
        cover_letter_content,
        output_folder,
        "2025_Cover_Letter.docx"  # Final filename
    )
    print(f"Final cover letter saved to: {final_cover_letter}")
    
    print("Job application process completed successfully!")